#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF (CV) -> texto -> JSON (GPT-4 preferente) -> Informe (GPT opcional) -> DOCX (formato Softtek)
"""

import os
import re
import json
import argparse
import datetime
from typing import Dict, Any, List
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from dotenv import load_dotenv, find_dotenv
load_dotenv(find_dotenv())

# ---------- imports opcionales/robustos ----------
def try_import(name: str):
    try:
        return __import__(name)
    except Exception:
        return None

PyPDF2 = try_import("PyPDF2")
pdfminer = try_import("pdfminer")
fitz = try_import("fitz")      # PyMuPDF
docx = try_import("docx")

# OpenAI SDK nuevo (openai>=1.x)
OpenAI = None
try:
    from openai import OpenAI as _OpenAI
    OpenAI = _OpenAI
except Exception:
    OpenAI = None

TECH_REGEX = re.compile(r"""
(java|kotlin|spring|spring\s+boot|spring\s+cloud|security|
kafka|grpc|rest|openapi|swagger|
mongo(db)?|couchbase|mysql|oracle|redis|
docker|kubernetes|gradle|maven|git(hub|lab)?|jenkins|ci/?cd|
aws|azure|ibm\s*cloud|lambda|sqs)
""", re.I | re.X)

# --- Catálogo ampliable y alias para visualización ---
TECH_CATALOG_DEFAULT = {
    # Lenguajes / runtimes
    "java","kotlin","scala","python","pyspark","c#","c++","c","go","php","ruby","typescript","javascript",".net",".net core",
    # Web / frameworks
    "spring","spring boot","spring cloud","hibernate","jpa","django","flask","fastapi","express","nestjs",
    "react","angular","vue","next.js","nuxt.js",
    # Build / CI
    "maven","gradle","npm","yarn","jenkins","github actions","gitlab ci","azure devops","argo cd","nexus","sonarqube",
    # VCS / herramientas
    "git","github","gitlab","bitbucket","intellij idea","jupyter notebooks","jira","confluence",
    # Contenedores / k8s
    "docker","kubernetes","helm","openshift","istio",
    # Mensajería
    "kafka","rabbitmq","activemq",
    # Data / BigData
    "apache spark","spark","pyspark","hadoop","hive","hdfs","airflow","nifi","databricks","snowflake",
    # DB / caches / search
    "postgresql","mysql","sql server","oracle","mongodb","cassandra","redis","elasticsearch",
    # Cloud AWS
    "aws","amazon s3","s3","athena","glue","emr","lambda","api gateway","rds","dynamodb","aurora","sns","sqs","kinesis","sagemaker","ec2",
    # Azure
    "azure","data factory","synapse","cosmos db",
    # GCP
    "gcp","bigquery","dataflow","dataproc","pub/sub",
    # BI / analítica
    "power bi","tableau","looker","superset",
    # Orquestadores / schedulers
    "control-m","oozie",
}

TECH_ALIASES = {
    "spark": "Apache Spark",
    "pyspark": "PySpark",
    "intellij": "IntelliJ IDEA",
    "intellij idea": "IntelliJ IDEA",
    "powerbi": "Power BI",
    "ms power bi": "Power BI",
    "s3": "Amazon S3",
    "aws s3": "Amazon S3",
    "athena": "Amazon Athena",
    "glue": "AWS Glue",
    "emr": "Amazon EMR",
    "lambda": "AWS Lambda",
    "rds": "Amazon RDS",
    "sns": "Amazon SNS",
    "sqs": "Amazon SQS",
    "kinesis": "Amazon Kinesis",
    "sagemaker": "Amazon SageMaker",
    "pub/sub": "Pub/Sub",
    "github": "GitHub",
    "gitlab": "GitLab",
    ".net": ".NET",
    ".net core": ".NET Core",
}

def _norm_token(s: str) -> str:
    s = str(s or "").strip()
    s = re.sub(r"\s+", " ", s)
    return s

def _canon_key(s: str) -> str:
    return re.sub(r"\s+", " ", s.lower().strip())

def _load_catalog_from_file(path: str) -> set:
    try:
        items = set()
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                t = line.strip()
                if t and not t.startswith("#"):
                    items.add(_canon_key(t))
        return items
    except Exception:
        return set()

def _mine_candidates_from_text(cv_text: str) -> List[str]:
    """
    Extrae candidatos técnicos desde el texto del CV:
    - tokens con '.', '+', '#', '/'
    - secuencias Capitalizadas de 1-3 palabras (Amazon S3, Apache Spark…)
    - keywords técnicas frecuentes en minúscula
    """
    text = re.sub(r"\s+", " ", cv_text or "")
    candidates = set()

    # C++, C#, .NET, Node.js, Next.js, etc.
    for m in re.finditer(r"\b(?:[A-Za-z]\w*(?:[./+#-]\w+)+)\b", text):
        candidates.add(m.group(0))

    # Secuencias Capitalizadas 1-3 palabras
    for m in re.finditer(r"\b([A-Z][a-zA-Z0-9]+(?:\s+[A-Z][a-zA-Z0-9]+){0,2})\b", text):
        cand = m.group(0)
        if len(cand) >= 3 and not re.match(r"^(Responsabilidades|Proyecto|Cliente|Empresa|Experiencia|Funciones)$", cand):
            candidates.add(cand)

    # keywords comunes
    kws = ["kafka","spark","pyspark","hadoop","hive","hdfs","airflow","nifi","docker","kubernetes","helm",
           "jenkins","terraform","ansible","git","bitbucket","gitlab","github","sonarqube",
           "postgresql","mysql","oracle","mongodb","redis","elasticsearch",
           "react","angular","vue","nestjs","django","flask","fastapi","express","spring","spring boot",".net",".net core",
           "aws","azure","gcp","s3","athena","glue","emr","lambda","rds","dynamodb","aurora","sns","sqs","kinesis",
           "bigquery","dataflow","dataproc","pub/sub","power bi","tableau","looker","control-m"]
    for kw in kws:
        if re.search(rf"\b{re.escape(kw)}\b", text, flags=re.I):
            candidates.add(kw)

    return list(candidates)

def deduce_tech_stack(cv_text: str, datos: Dict[str, Any], catalog_path: str = "assets/tech_catalog.txt", limit: int = 60) -> List[str]:
    """
    Combina tecnologías del JSON + minería de texto del CV.
    Filtra contra catálogo (archivo + default). Aplica alias y dedupe con orden estable.
    """
    # 1) tokens desde JSON
    j_tokens = []
    j_tokens += [sanitize(t) for t in (datos.get("skills") or [])]
    for e in (datos.get("experiencia") or []):
        j_tokens += [sanitize(t) for t in (e.get("tecnologias") or [])]
    for t in (datos.get("habilidades", {}).get("herramientas") or []):
        j_tokens.append(sanitize(t))

    # 2) tokens desde texto
    mined = _mine_candidates_from_text(cv_text)

    # 3) catálogo
    cat_file = _load_catalog_from_file(catalog_path)
    catalog = set(TECH_CATALOG_DEFAULT) | cat_file

    # 4) normalización, alias y dedupe
    seen = set()
    out = []

    def push(tok: str):
        tok = _norm_token(tok)
        if not tok:
            return
        visual = TECH_ALIASES.get(_canon_key(tok), TECH_ALIASES.get(tok.lower(), tok))
        k2 = _canon_key(visual)
        if k2 in seen:
            return
        seen.add(k2)
        out.append(visual)

    def admissible(tok: str) -> bool:
        k = _canon_key(tok)
        if k in catalog:
            return True
        # coincidencia parcial laxa para no perder variantes (p. ej. "Amazon S3" vs "S3")
        for c in catalog:
            if k == c:
                return True
            if len(c) >= 3 and (k in c or c in k):
                return True
        return False

    for t in j_tokens:
        if t and admissible(t):
            push(t)
    for t in mined:
        if t and admissible(t):
            push(t)

    return out[:limit]

# ---------- utilidades ----------
def extract_text_from_pdf(pdf_path: str) -> str:
    """Cascada: PyPDF2 -> pdfminer -> PyMuPDF."""
    if PyPDF2:
        try:
            parts = []
            with open(pdf_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for p in reader.pages:
                    parts.append(p.extract_text() or "")
            txt = "\n".join(parts)
            if txt.strip():
                return txt
        except Exception:
            pass
    if pdfminer:
        try:
            from pdfminer.high_level import extract_text as pm_extract
            txt = pm_extract(pdf_path)
            if txt.strip():
                return txt
        except Exception:
            pass
    if fitz:
        try:
            doc = fitz.open(pdf_path)
            parts = []
            for page in doc:
                # Cada bloque = (x0, y0, x1, y1, text, block_no, ...)
                blocks = page.get_text("blocks")
                blocks.sort(key=lambda b: (b[1], b[0]))  # orden por y, luego x
                parts.extend(b[4] for b in blocks if isinstance(b[4], str) and b[4].strip())
            txt = "\n".join(parts)
            if txt.strip():
                return txt
        except Exception:
            pass
    raise RuntimeError("No fue posible extraer texto del PDF. Instala PyPDF2 o pdfminer.six o PyMuPDF.")

def solo_tecnologias(tokens):
    clean = []
    for t in tokens or []:
        t2 = sanitize(t)
        if TECH_REGEX.search(t2):
            clean.append(t2)
    # dedupe y limita
    return list(dict.fromkeys(clean))[:40]

def _safe_json_loads(maybe_json: str) -> Dict[str, Any]:
    """Decodifica JSON incluso con 'ruido' alrededor."""
    try:
        return json.loads(maybe_json)
    except Exception:
        pass
    m = re.search(r"(\{.*\}|\[.*\])", maybe_json, re.S)
    if m:
        try:
            return json.loads(m.group(1))
        except Exception:
            pass
    fixed = re.sub(r",\s*([}\]])", r"\1", maybe_json).replace("'", '"')
    return json.loads(fixed)

def sanitize(x: object) -> str:
    """Sin None, sin [] ni ""."""
    txt = str(x or "").strip()
    return re.sub(r'[\[\]"]+', "", txt)

def _split_in_two_paragraphs(texto: str):
    """Convierte un resumen largo en 1–2 párrafos legibles."""
    t = re.sub(r"\s+", " ", (texto or "").strip())
    if not t:
        return []
    # si ya hay párrafos separados por líneas en blanco, respétalos
    parts = [p.strip() for p in re.split(r"\n\s*\n+", t) if p.strip()]
    if len(parts) >= 2:
        return parts[:2]
    # si no, corta por oraciones más o menos a la mitad
    sents = re.split(r"(?<=[.!?…])\s+", t)
    if len(sents) <= 2:
        return [t]
    mid = max(1, len(sents)//2)
    return [" ".join(sents[:mid]).strip(), " ".join(sents[mid:]).strip()]

def _cell_label(cell, text, underline=True):
    """Etiqueta con formato configurable dentro de una celda."""
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.underline = bool(underline)

def _cell_paragraphs(cell, paragraphs):
    """Escribe varios párrafos (líneas) en una celda."""
    cell.text = ""
    if not paragraphs:
        return
    p = cell.paragraphs[0]
    p.add_run(paragraphs[0])
    for para in paragraphs[1:]:
        p = cell.add_paragraph()
        p.add_run(para)

def hide_table_borders(table):
    """Quita todos los bordes visibles de la tabla (exterior e interiores)."""
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        tblBorders.append(el)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)

def set_cell_margins(table, top=60, left=80, bottom=60, right=80):
    """Margen interno de celdas en DXA (1/20 pt). 80≈4pt≈1.4mm."""
    for cell in table._cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        old = tcPr.find(qn('w:tcMar'))
        if old is not None:
            tcPr.remove(old)
        tcPr.append(tcMar)

def postprocesar_json(cv_text: str, datos: Dict[str, Any]) -> Dict[str, Any]:
    """Limpia 'null' literales, rellena roles vacíos y amplía descripciones breves usando experiencia_funcional si existe."""
    def is_nullish(v):
        return v is None or (isinstance(v, str) and v.strip().lower() in {"null","none","n/a","na","-","—","--",""} )

    # 1) Educación: evita mostrar "null" como literal
    for ed in datos.get("educacion", []) or []:
        if isinstance(ed.get("titulacion"), str) and ed["titulacion"].strip().lower() in {"null","none","n/a","na","-","—","--"}:
            ed["titulacion"] = None

    # 2) Experiencia: role fallback + descripción más larga
    hab = (datos.get("habilidades") or {})
    rol_fallback = hab.get("perfil_principal") or "Back-End Engineer"
    for e in datos.get("experiencia", []) or []:
        if is_nullish(e.get("rol")):
            e["rol"] = rol_fallback

        # si hay experiencia_funcional y descripcion es corta, usar la funcional
        desc = (e.get("descripcion") or "").strip()
        funcional = (e.get("experiencia_funcional") or "").strip() if isinstance(e.get("experiencia_funcional"), str) else ""
        if len(desc) < 220 and len(funcional) > len(desc):
            e["descripcion"] = funcional

        # si sigue muy corta, intenta unir líneas (quita cortes y dobles espacios)
        if e.get("descripcion"):
            clean = re.sub(r"\s+", " ", e["descripcion"]).strip()
            e["descripcion"] = clean

        # dedupe tecnologías
        techs = [t.strip() for t in (e.get("tecnologias") or []) if t and str(t).strip()]
        e["tecnologias"] = list(dict.fromkeys(techs))

    # 3) Habilidades: dedupe herramientas y limita
    herr = [h.strip() for h in (hab.get("herramientas") or []) if h and str(h).strip()]
    datos.setdefault("habilidades", {})
    datos["habilidades"]["herramientas"] = list(dict.fromkeys(herr))[:60]

    return datos


# ---------- JSON del CV con GPT-4 (preferente) ----------
def gpt_cv_text_to_json(cv_text: str) -> Dict[str, Any]:
    """Usa GPT-4(o) para producir JSON canónico. Requiere OPENAI_API_KEY y SDK nuevo."""
    if not OpenAI or not os.getenv("OPENAI_API_KEY"):
        raise RuntimeError("No hay OpenAI SDK o API key.")

    client = OpenAI()

    schema_hint = {
        "type": "object",
        "properties": {
            "nombre": {"type": "string"},
            "email": {"type": "string"},
            "telefono": {"type": ["string", "null"]},
            "ubicacion": {"type": ["string", "null"]},
            "titulo": {"type": ["string", "null"]},
            "links": {"type": "array", "items": {"type": "string"}},
            "idiomas": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "idioma": {"type": "string"},
                        "nivel": {"type": ["string", "null"]}
                    },
                    "required": ["idioma"]
                }
            },
            "habilidades": {
                "type": "object",
                "properties": {
                    "perfil_principal": {"type": ["string", "null"]},
                    "perfil_secundario": {"type": ["string", "null"]},
                    "herramientas": {"type": "array", "items": {"type": "string"}}
                },
                "required": ["perfil_principal", "herramientas"]
            },
            "skills": {"type": "array", "items": {"type": "string"}},
            "experiencia": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "empresa": {"type": "string"},
                        "rol": {"type": ["string", "null"]},
                        "ubicacion": {"type": ["string", "null"]},
                        "desde": {"type": ["string", "null"]},
                        "hasta": {"type": ["string", "null"]},
                        "descripcion": {"type": ["string", "null"]},
                        "tecnologias": {"type": "array", "items": {"type": "string"}}
                    },
                    "required": ["empresa"]
                }
            },
            "educacion": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "centro": {"type": "string"},
                        "titulacion": {"type": ["string", "null"]},
                        "fecha": {"type": ["string", "null"]}
                    },
                    "required": ["centro"]
                }
            },
            "certificaciones": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "nombre": {"type": "string"},
                        "emisor": {"type": ["string", "null"]},
                        "fecha-inicio": {"type": ["string", "null"]},
                        "fecha-fin": {"type": ["string", "null"]}
                    },
                    "required": ["nombre"]
                }
            },
            "cursos": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                    "nombre": {"type": "string"},
                    "entidad": {"type": ["string", "null"]},
                    "fecha": {"type": ["string", "null"]}
                    },
                    "required": ["nombre"]
                }
            },
            "resumen": {"type": ["string", "null"]}
        },
        "required": ["nombre", "experiencia"]
    }

    system = (
        "Eres un extractor experto de CVs. Devuelve SOLO JSON válido UTF-8, sin texto extra. "
        "No inventes datos; si faltan usa null o []. "
        "Normaliza meses (Ene/Jan… Dic/Dec) y fechas a 'YYYY' o 'YYYY-MM'. "
        "Si un puesto sigue activo, 'hasta' = 'Hoy'. "
        "Para cada experiencia, COPIA (en español) el contenido del apartado de experiencia del CV: "
        "usa hasta 2–3 PÁRRAFOS completos combinando texto corrido y bullets si existen; "
        "si hay varios proyectos en la misma empresa, menciónalos de forma concisa. "
        "NO inventes; si no hay contenido suficiente, deja 1 párrafo. "
        "Además, al estraer la experiencia, tranformala en tercera persona de manera que sea una descripción de un agente externo de la experiencia del candidato."
        "Usa un tono netro y profesional para las experiencias funcionales, con sintagmas nominales y verbos en gerundio. Con foco técnico y contexto de negocio"
        "Extrae 'tecnologias' como lista (Java, Spring Boot, Kafka, MongoDB, Kubernetes, AWS, etc.). "
        "En 'educacion' asegurate de extraer {centro, titulacion, fecha}."
        "En 'habilidades' detecta: perfil_principal, perfil_secundario  ( el segundo que más aparezca ) y lista de herramientas técnicas. "
        "Si hay CERTIFICACIONES, devuelve {nombre, emisor, fecha-inicio, fecha-fin}."
        "Si existe un bloque CURSOS/COURSES, extrae cada curso con {nombre, entidad, fecha}."
    )

    user = f"Esquema guía: {json.dumps(schema_hint, ensure_ascii=False)}\nCV (texto):\n{cv_text}"

    res = client.chat.completions.create(
        model=os.getenv("OPENAI_MODEL_JSON", "gpt-4o-mini"),
        messages=[{"role": "system", "content": system},
                  {"role": "user", "content": user}],
        temperature=0.0,
        max_tokens=2000
    )
    datos = _safe_json_loads(res.choices[0].message.content)

    # defaults
    datos.setdefault("links", [])
    datos.setdefault("idiomas", [])
    datos.setdefault("skills", [])
    datos.setdefault("experiencia", [])
    datos.setdefault("educacion", [])
    datos.setdefault("cursos", [])
    datos["extraido_en"] = datetime.datetime.now().isoformat()
    return datos

# ---------- Heurística fallback (si no hay API) ----------
def parse_cv_text_to_json(text: str) -> Dict[str, Any]:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    big = "\n".join(lines)

    # ---- helpers locales ----
    def _slice_between(big_text: str, start_kw: str, end_kws: List[str]) -> str:
        s = re.search(rf"{start_kw}\s*[:\n]", big_text, re.I)
        if not s:
            return ""
        start = s.end()
        end = len(big_text)
        for ek in end_kws:
            m = re.search(rf"\n\s*{ek}\s*[:\n]", big_text[start:], re.I)
            if m:
                end = start + m.start()
                break
        return big_text[start:end].strip()

    # ---- nombre ----
    name = None
    for l in lines[:10]:
        if re.match(r"^[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+){1,3}$", l):
            name = l
            break
    if not name:
        name = "Candidato/a"

    # ---- email ----
    email = ""
    m = re.search(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", big)
    if m:
        email = m.group(0)

    # ---- título (líneas iniciales) ----
    title = ""
    for l in lines[:30]:
        if re.search(r"(backend|engineer|ingenier|desarrollador|java|kotlin)", l, re.I):
            title = l
            break

    # ---- skills rápidas (visibles) ----
    tech_candidates = re.findall(r"\b(Java|Kotlin|Spring|Kafka|MongoDB|Redis|Docker|Kubernetes|AWS|Azure)\b", big, re.I)
    skills = sorted({t.title() for t in tech_candidates})

    # ---- objeto base ----
    data: Dict[str, Any] = {
        "nombre": name,
        "email": email,
        "ubicacion": "",
        "titulo": title,
        "links": sorted(set(re.findall(r"https?://[^\s)]+", big))),
        "idiomas": [],
        "experiencia": [],
        "educacion": [],
        "skills": skills,
        "extraido_en": datetime.datetime.now().isoformat()
    }

    # ----- CURSOS (heurística: solo si hay bloque) -----
    cursos_raw = _slice_between(
        big,
        r"Cursos|Courses|Formación complementaria",
        ["Certificaciones", "Certifications", "Educación", "Formación", "Experience", "Experiencia"]
    )
    cursos: List[Dict[str, Any]] = []
    if cursos_raw:
        for line in [l.strip("•- \t") for l in cursos_raw.splitlines() if l.strip()]:
            # "Curso – Entidad (Fecha)"  /  "Curso - Entidad (Fecha)"  /  "Curso"
            m = re.match(r"(.+?)\s*[–-]\s*(.+?)\s*(\(([^)]+)\))?$", line)
            if m:
                nombre = m.group(1).strip()
                entidad = (m.group(2) or "").strip() or None
                fecha = (m.group(4) or "").strip() or None
            else:
                nombre, entidad, fecha = line, None, None
            if nombre:
                cursos.append({"nombre": nombre, "entidad": entidad, "fecha": fecha})
    data["cursos"] = cursos  # quedará [] si no hay bloque

    return data



# ---------- Informe con GPT (opcional) ----------
def redactar_informe_con_gpt(rol: str, datos: Dict[str, Any]) -> str:
    if not OpenAI or not os.getenv("OPENAI_API_KEY"):
        nombre = datos.get("nombre", "el/la candidato/a")
        titulo = sanitize(datos.get("titulo", ""))
        idi = ", ".join([f"{i['idioma']} ({i.get('nivel','')})" for i in datos.get("idiomas", [])]) or "No indicado"
        skills = ", ".join(datos.get("skills", [])[:12]) or "Java, Spring Boot, Microservicios"
        bullets = []
        for e in datos.get("experiencia", [])[:5]:
            bullets.append(f"• {sanitize(e.get('empresa'))} — {sanitize(e.get('rol'))} ({sanitize(e.get('desde'))} - {sanitize(e.get('hasta'))})")
        exp_txt = "\n".join(bullets) or "• Experiencia no detectada automáticamente."
        return f"""Resumen del perfil
{nombre} — {titulo}

Ajuste al rol y fortalezas
Encaja con el rol (Java, Spring, Kafka, NoSQL, TDD, CI/CD, e-commerce). Fortalezas: calidad, DDD, cloud.

Experiencia relevante
{exp_txt}

Competencias técnicas
{skills}

Idiomas y disponibilidad
Idiomas: {idi}. Disponibilidad estimada: 2-4 semanas.
"""
    client = OpenAI()
    prompt = f"""Redacta un informe (1-2 páginas) del candidato para el rol:
{rol}

Datos estructurados:
{json.dumps(datos, ensure_ascii=False, indent=2)}

Estructura:
- Resumen del perfil
- Ajuste al rol y fortalezas
- Experiencia relevante (bullets)
- Competencias técnicas
- Idiomas y disponibilidad
- Riesgos/áreas de mejora y mitigaciones
"""
    res = client.chat.completions.create(
        model=os.getenv("OPENAI_MODEL_REPORT", "gpt-4o-mini"),
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    return res.choices[0].message.content.strip()

# ---------- DOCX formato Softtek ----------
def generar_docx_softtek(datos: Dict[str, Any], salida_docx: str, logo_path: str = None, cv_text_raw: str = ""):
    # --- fallback si no hay python-docx ---
    if not docx:
        with open(os.path.splitext(salida_docx)[0] + ".txt", "w", encoding="utf-8") as f:
            f.write("Informe (texto):\n\n")
            f.write(json.dumps(datos, ensure_ascii=False, indent=2))
        return

    # --- helpers locales ---
    import os, re, datetime as _dt
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # usa sanitize global si existe; si no, crea uno local
    try:
        s = sanitize
    except NameError:
        def s(x):
            txt = str(x or "").strip()
            return re.sub(r'[\[\]"]+', "", txt)

    MESES_ES = ["Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]

    def _fmt_fecha_es(fecha):
        if not fecha:
            return ""
        f = str(fecha).strip()
        fl = f.lower()
        if fl.startswith(("hoy", "actual")) or "presente" in fl:
            return "Actualidad"
        m = re.match(r"(\d{4})-(\d{2})", f)  # YYYY-MM
        if m:
            y, mm = int(m.group(1)), int(m.group(2))
            if 1 <= mm <= 12:
                return f"{MESES_ES[mm-1]} {y}"
        m = re.search(r"(\d{4})", f)  # YYYY
        return m.group(1) if m else f

    def _parse_fecha(sv):
        if not sv:
            return None
        s2 = str(sv).strip().lower()
        if s2.startswith(("hoy", "actual")) or "presente" in s2:
            t = _dt.date.today()
            return _dt.date(t.year, t.month, 1)
        m = re.match(r"(\d{4})-(\d{2})", s2)  # YYYY-MM
        if m:
            return _dt.date(int(m.group(1)), int(m.group(2)), 1)
        m = re.search(r"(\d{4})", s2)  # YYYY
        if m:
            return _dt.date(int(m.group(1)), 1, 1)
        return None

    def _duracion(desde, hasta):
        d0 = _parse_fecha(desde)
        d1 = _parse_fecha(hasta) or _parse_fecha("hoy")
        if not d0 or not d1:
            return ""
        months = (d1.year - d0.year) * 12 + (d1.month - d0.month)
        months = max(0, months)
        y, m = divmod(months, 12)
        if y and m:
            return f"{y} año{'s' if y != 1 else ''} y {m} mes{'es' if m != 1 else ''}"
        if y:
            return f"{y} año{'s' if y != 1 else ''}"
        if m:
            return f"{m} mes{'es' if m != 1 else ''}"
        return "menos de 1 mes"

    def _anios_totales(exp_list):
        """Años totales aproximados: desde el inicio más antiguo hasta el fin más reciente."""
        fechas = []
        for e in exp_list or []:
            d0 = _parse_fecha(e.get("desde"))
            d1 = _parse_fecha(e.get("hasta")) or _parse_fecha("hoy")
            if d0 and d1:
                fechas.append((d0, d1))
        if not fechas:
            return 0
        start = min(f[0] for f in fechas)
        end = max(f[1] for f in fechas)
        months = (end.year - start.year) * 12 + (end.month - start.month)
        return max(1, round(months / 12))  # al menos 1 año si hay experiencia

    TECH_REGEX = re.compile(r"""
        (java|kotlin|spring(\s+boot|\s+cloud)?|security|
         kafka|grpc|rest|openapi|swagger|
         mongo(db)?|couchbase|mysql|oracle|redis|
         docker|kubernetes|gradle|maven|git(hub|lab)?|jenkins|ci/?cd|
         aws|azure|gcp|lambda|sqs|sns|cloud)
    """, re.I | re.X)

    def solo_tecnologias(tokens):
        clean = []
        for t in tokens or []:
            t2 = s(t)
            if TECH_REGEX.search(t2):
                clean.append(t2)
        return list(dict.fromkeys([t for t in clean if t]))[:40]

    MAP_NIV = {
        "native": "Nativo", "bilingual": "Bilingüe",
        "fluent": "C1", "advanced": "C1",
        "upper-intermediate": "B2", "intermediate": "B2",
        "basic": "A2"
    }

    # --- documento ---
    document = Document()

    # márgenes
    for sec in document.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # Encabezado SOLO en la primera página
    sec = document.sections[0]
    try:
        # Activa encabezado/pie distintos para la primera página
        sec.different_first_page_header_footer = True

        # Escribe el logo o texto corporativo en el header de la primera página
        hdr = sec.first_page_header
        p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        p.alignment = WD_TABLE_ALIGNMENT.LEFT  # mismo enum value para alinear izda
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()

        if logo_path and os.path.exists(logo_path):
            try:
                run.add_picture(logo_path, width=Cm(4.0))
            except Exception:
                from docx.shared import RGBColor
                r = p.add_run("Softtek")
                r.bold = True
                r.font.size = Pt(20)
                r.font.color.rgb = RGBColor(0, 100, 177)
        else:
            from docx.shared import RGBColor
            r = p.add_run("Softtek")
            r.bold = True
            r.font.size = Pt(20)
            r.font.color.rgb = RGBColor(0, 100, 177)

        # Deja el header "normal" vacío
    except AttributeError:
        # Fallback: insertar logo/brand en el cuerpo
        p = document.add_paragraph()
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if logo_path and os.path.exists(logo_path):
            run = p.add_run()
            try:
                run.add_picture(logo_path, width=Cm(4.0))
            except Exception:
                from docx.shared import RGBColor
                r = p.add_run("Softtek")
                r.bold = True
                r.font.size = Pt(20)
                r.font.color.rgb = RGBColor(0, 100, 177)
        else:
            from docx.shared import RGBColor
            r = p.add_run("Softtek")
            r.bold = True
            r.font.size = Pt(20)
            r.font.color.rgb = RGBColor(0, 100, 177)

    document.add_paragraph()  # espacio

    # Título
    t = document.add_paragraph()
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(f"{s(datos.get('nombre'))}")
    tr.bold = True
    tr.font.size = Pt(14)

    def add_title(txt):
        p = document.add_paragraph()
        r = p.add_run(txt.upper() + ":")
        r.bold = True
        r.underline = True

    document.add_paragraph()

    # PROFILE (primer bullet con título + años, luego roles con duración)
    add_title("PROFILE")
    exp_list = datos.get("experiencia") or []
    titulo = s(datos.get("titulo") or "Profesional TI")
    anios = _anios_totales(exp_list)

    from docx.shared import Cm  # asegúrate de tener esta import arriba

    def add_bullet(texto: str, level: int = 0):
        """Crea un bullet y le aplica indentación tipo tab según level."""
        p = document.add_paragraph(style="List Bullet")
        # Indentación visual (tabulado)
        if level > 0:
            p.paragraph_format.left_indent = Cm(0.75 * level)
            p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(texto)

    # Un bullet por cada experiencia encontrada (sin límite)
    for e in exp_list:
        rol = s(e.get("rol") or "Back-End Engineer")
        emp = s(e.get("empresa") or "")
        dur = _duracion(e.get("desde"), e.get("hasta"))
        add_bullet(f"{rol} – {emp} ({dur})", level=1)

    document.add_paragraph()

    # --- EXPERIENCIA PROFESIONAL (TABLAS SIN BORDES) ---
    add_title("EXPERIENCIA PROFESIONAL")
    for e in (datos.get("experiencia") or []):
        empresa = s(e.get("empresa"))
        ubic = s(e.get("ubicacion") or "")
        rol = s(e.get("rol") or (datos.get("habilidades") or {}).get("perfil_principal"))
        desde_es = _fmt_fecha_es(s(e.get("desde")))
        hasta_es = _fmt_fecha_es(s(e.get("hasta"))) or "Actualidad"

        techs = [s(t) for t in (e.get("tecnologias") or []) if s(t)]
        techs = list(dict.fromkeys(techs))  # dedupe

        # si tienes función que parte a 2 párrafos:
        resumen_paras = _split_in_two_paragraphs(e.get("descripcion") or "")

        # crea tabla 3x5
        table = document.add_table(rows=3, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True

        # quitar bordes
        hide_table_borders(table)
        # márgenes internos
        set_cell_margins(table, top=60, left=80, bottom=60, right=80)

        # Fila 1
        c = table.rows[0].cells
        c[0].text = empresa
        c[1].text = ubic
        c[2].text = rol
        c[3].text = f"Desde {desde_es}"
        c[4].text = f"Hasta: {hasta_es}"
        for i in range(5):
            run = c[i].paragraphs[0].runs[0] if c[i].paragraphs[0].runs else c[i].paragraphs[0].add_run(c[i].text)
            run.bold = True

        # Fila 2 (Experiencia funcional)
        _cell_label(table.rows[1].cells[0], "Experiencia funcional:", underline=False)
        cell_desc = table.rows[1].cells[1]
        cell_desc.merge(table.rows[1].cells[4])  # colspan 4
        _cell_paragraphs(cell_desc, resumen_paras or [s(e.get("descripcion") or "")])

        # Fila 3 (Herramientas)
        _cell_label(table.rows[2].cells[0], "Herramientas:", underline=True)
        cell_tools = table.rows[2].cells[1]
        cell_tools.merge(table.rows[2].cells[4])   # colspan 4
        cell_tools.text = ", ".join(techs) if techs else "—"

        document.add_paragraph()  # espacio
        document.add_paragraph()

    # FORMACIÓN
    add_title("FORMACIÓN ACADÉMICA")
    for ed in (datos.get("educacion") or []):
        p = document.add_paragraph()
        tit = s(ed.get("titulacion") or "")
        if tit.lower() in ("null","none","n/a","na","-","—","--"):
            tit = ""
        cen = s(ed.get("centro"))
        fec = s(ed.get("fecha"))
        if tit and cen:
            p.add_run(f"{tit}. {cen}  ")
        else:
            p.add_run(f"{tit or cen}  ")
        if fec:
            p.add_run(fec)

    document.add_paragraph()

    # CURSOS (si existen)
    if datos.get("cursos"):
        add_title("CURSOS")
        for c in (datos.get("cursos") or []):
            nombre = s(c.get("nombre"))
            entidad = s(c.get("entidad"))
            fecha = s(c.get("fecha"))
            line = nombre
            if entidad: line += f". {entidad}"
            if fecha:   line += f"  {fecha}"
            p = document.add_paragraph()
            p.add_run(line)

        document.add_paragraph()

    # CERTIFICACIONES (si existen)
    certs = datos.get("certificaciones")
    if certs:
        add_title("CERTIFICACIONES")
        for c in certs:
            p = document.add_paragraph()
            name = s(c.get("nombre") or c.get("emisor"))
            p.add_run(name)
            if c.get("fecha"):
                p.add_run(f" — {s(c['fecha'])}")

    document.add_paragraph()

    # IDIOMAS
    add_title("IDIOMAS")
    idiomas = datos.get("idiomas") or []
    if not idiomas:
        p = document.add_paragraph()
        p.add_run("—")
    else:
        for i in idiomas:
            idioma = s(i.get("idioma")).capitalize()
            nivel = s(i.get("nivel") or "")
            nivel_es = MAP_NIV.get(nivel.lower(), nivel) if nivel else "—"
            p = document.add_paragraph()
            p.add_run(f"{idioma}: {nivel_es}")

    document.add_paragraph()

    # HABILIDADES TÉCNICAS (perfil principal/secundario + herramientas)
    add_title("HABILIDADES TÉCNICAS")

    hab = datos.get("habilidades") or {}
    pp = s(hab.get("perfil_principal") or datos.get("perfil_principal") or "")
    ps = s(hab.get("perfil_secundario") or datos.get("perfil_secundario") or "")

    if pp:
        p = document.add_paragraph()
        r = p.add_run("Perfil principal: "); r.bold = True
        p.add_run(pp)

    if ps:
        p = document.add_paragraph()
        r = p.add_run("Perfil Secundario: "); r.bold = True
        p.add_run(ps if ps.endswith(".") else (ps + "."))

    # --- NUEVO: deducir herramientas desde JSON + texto del CV + catálogo/alias ---
    tools = deduce_tech_stack(cv_text_raw or "", datos, catalog_path="assets/tech_catalog.txt", limit=60)

    p = document.add_paragraph()
    r = p.add_run("Herramientas: "); r.bold = True
    p.add_run(", ".join(tools) if tools else "—")

    # guardar
    os.makedirs(os.path.dirname(salida_docx) or ".", exist_ok=True)
    document.save(salida_docx)


# ---------- CLI ----------
def main():
    ap = argparse.ArgumentParser(description="CV PDF -> JSON (GPT-4) -> Informe -> DOCX Softtek")
    ap.add_argument("--cv_pdf", required=True)
    ap.add_argument("--rol_txt", required=True)
    ap.add_argument("--salida_json", required=True)
    ap.add_argument("--salida_docx", required=True)
    ap.add_argument("--logo_path", default=None)
    args = ap.parse_args()

    texto = extract_text_from_pdf(args.cv_pdf)

    try:
        datos = gpt_cv_text_to_json(texto)
    except Exception:
        datos = parse_cv_text_to_json(texto)

    datos = postprocesar_json(texto, datos)

    rol = ""
    if args.rol_txt and os.path.exists(args.rol_txt):
        with open(args.rol_txt, "r", encoding="utf-8") as f:
            rol = f.read().strip()

    informe = redactar_informe_con_gpt(rol, datos)

    os.makedirs(os.path.dirname(args.salida_json) or ".", exist_ok=True)
    with open(args.salida_json, "w", encoding="utf-8") as f:
        json.dump(datos, f, ensure_ascii=False, indent=2)

    # <-- pásale el texto del CV para deducir herramientas
    generar_docx_softtek(datos, args.salida_docx, logo_path=args.logo_path, cv_text_raw=texto)

    print("✔ JSON guardado en:", args.salida_json)
    print("✔ DOCX guardado en:", args.salida_docx)

if __name__ == "__main__":
    main()
